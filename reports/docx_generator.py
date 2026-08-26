from pathlib import Path
from docx import Document
from docx.shared import Inches
from reports.pdf_generator import DISCLAIMER


def generate_docx(result: dict, assets: dict, destination: Path) -> Path:
    doc=Document(); doc.add_heading("NeuroVisionAssisted Brain MRI Analysis",0); doc.add_paragraph(f"Study ID: {result['study_id']}"); doc.add_paragraph(DISCLAIMER)
    doc.add_heading("Study Information",1); doc.add_paragraph(f"Detected segmented regions: {result['tumor_count']}")
    for label in ("original","annotated"):
        if assets.get(label): doc.add_picture(str(assets[label]),width=Inches(5.5))
    doc.add_heading("Tumor Summary",1); table=doc.add_table(rows=1, cols=4); table.style="Light Shading Accent 1"
    for cell,text in zip(table.rows[0].cells,["ID","Area (px)","Max diameter (px)","Centroid"]): cell.text=text
    for t in result["tumors"]:
        cells=table.add_row().cells
        for cell,text in zip(cells,[t["tumor_id"],str(t["area_pixels"]),f"{t['max_diameter_pixels']:.1f}",str(t["centroid"])]): cell.text=text
    if result["pairwise_analysis"]:
        doc.add_heading("Tumor-to-Tumor Distances",1)
        for p in result["pairwise_analysis"]: doc.add_paragraph(f"{p['tumor_a']}–{p['tumor_b']}: centroid {p['centroid_distance_pixels']:.1f}px; boundary {p['boundary_distance_pixels']:.1f}px; {p['relative_position']}")
    doc.add_heading("Segmented Regions Analysis", 1)
    for i, t in enumerate(result["tumors"]):
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(3.5)
        
        cell_img = table.cell(0, 0)
        cell_text = table.cell(0, 1)
        
        if i < len(assets.get("crops", [])):
            cell_img.paragraphs[0].add_run().add_picture(str(assets["crops"][i]), width=Inches(2.2))
        else:
            cell_img.paragraphs[0].text = "[Crop not available]"
            
        cell_text.paragraphs[0].text = f"Segmented Region {t['tumor_id']}"
        cell_text.paragraphs[0].style = doc.styles["Heading 2"]
        
        p = cell_text.add_paragraph()
        p.add_run(f"Bounding box: {t['bbox']}\n")
        p.add_run(f"Equivalent radius: {t['equivalent_radius_pixels']:.1f} pixels")
        doc.add_paragraph()
    doc.add_heading("Model and Limitations",1); doc.add_paragraph(f"Model: {result['model']['name']} v{result['model']['version']}. Segmentation metrics are placeholders pending validation.")
    doc.add_paragraph("Supportive health-information and specialist-information placeholders: consult a qualified medical professional for interpretation and next steps.")
    doc.save(destination); return destination
